"""
Tools the agent can use:
    1. web_search    - DuckDuckGo search
    2. docx_writer   - saves report as .docx
    3. txt_writer    - saves report as .txt (fallback)
    4. read_file     - reads saved files
    5. merge_results - merges several result blocks
"""

import os
from datetime import datetime
from tavily import TavilyClient
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def web_search(query: str, max_results: int = 5) -> str:
    """Search the web using Tavily."""
    try:
        api_key = os.environ.get("TAVILY_API_KEY")
        if not api_key:
            return "Search error: TAVILY_API_KEY not set."
        client = TavilyClient(api_key=api_key)
        response = client.search(query, max_results=max_results)
        results = response.get("results", [])
        if not results:
            return "No results found."
        formatted = []
        for i, r in enumerate(results, 1):
            snippet = r.get("content", "")[:200]  # FIX 2: truncate to 200 chars
            formatted.append(f"[{i}] {r.get('title', '')}\n{snippet}\nSource: {r.get('url', '')}")
        return "\n\n".join(formatted)
    except Exception as e:
        return f"Search error: {str(e)}"


def docx_writer(filename: str, title: str, summary: str, sections: list, sources: list) -> str:
    """Save report as a .docx file."""
    try:
        os.makedirs("reports", exist_ok=True)
        timestamp = datetime.now().strftime("%m%d%Y_%H%M%S")
        safe_name = filename.replace(" ", "_").replace("/", "_")
        filepath = f"reports/{safe_name}_{timestamp}.docx"

        doc = Document()

        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Topic: {filename}")
        doc.add_paragraph("-" * 60)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

        for section in sections:
            doc.add_heading(section.get("heading", ""), level=2)
            doc.add_paragraph(section.get("content", ""))

        doc.add_heading("Sources", level=1)
        for i, source in enumerate(sources, 1):
            doc.add_paragraph(f"[{i}] {source}")

        doc.save(filepath)
        return f"Report saved to: {filepath}"
    except Exception as e:
        return f"Save error: {str(e)}"


def txt_writer(filename: str, title: str, summary: str, sections: list, sources: list) -> str:
    """Save report as a plain .txt file."""
    try:
        os.makedirs("reports", exist_ok=True)
        timestamp = datetime.now().strftime("%m%d%Y_%H%M%S")
        safe_name = filename.replace(" ", "_").replace("/", "_")
        filepath = f"reports/{safe_name}_{timestamp}.txt"

        lines = [
            title,
            "=" * 60,
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Topic: {filename}",
            "-" * 60,
            "",
            "SUMMARY",
            "-" * 60,
            summary,
            "",
        ]
        for section in sections:
            lines.append(section.get("heading", "").upper())
            lines.append("-" * 40)
            lines.append(section.get("content", ""))
            lines.append("")

        lines.append("SOURCES")
        lines.append("-" * 60)
        for i, source in enumerate(sources, 1):
            lines.append(f"[{i}] {source}")

        with open(filepath, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

        return f"Report saved to: {filepath}"
    except Exception as e:
        return f"Save error: {str(e)}"


def read_file(filepath: str) -> str:
    """Read a saved file."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return f"File not found: {filepath}"
    except Exception as e:
        return f"Read error: {str(e)}"


def merge_results(results: list) -> str:
    """Merge multiple search result blocks into one formatted string."""
    merged = []
    for i, result in enumerate(results, 1):
        merged.append(f"=== Search Result Block {i} ===\n{result}")
    return "\n\n".join(merged)


_WRITER_TOOL = {
    "docx": {
        "name": "docx_writer",
        "description": "Save the final report as a .docx Word file when research is complete.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Report filename (no extension)"},
                "title": {"type": "string", "description": "Full report title"},
                "summary": {"type": "string", "description": "Executive summary paragraph"},
                "sections": {
                    "type": "array",
                    "items": {"type": "object", "properties": {"heading": {"type": "string"}, "content": {"type": "string"}}},
                    "description": "List of sections with heading and content"
                },
                "sources": {"type": "array", "items": {"type": "string"}, "description": "List of source URLs or citations"}
            },
            "required": ["filename", "title", "summary", "sections", "sources"]
        }
    },
    "txt": {
        "name": "txt_writer",
        "description": "Save the final report as a plain .txt file when research is complete.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Report filename (no extension)"},
                "title": {"type": "string", "description": "Full report title"},
                "summary": {"type": "string", "description": "Executive summary paragraph"},
                "sections": {
                    "type": "array",
                    "items": {"type": "object", "properties": {"heading": {"type": "string"}, "content": {"type": "string"}}},
                    "description": "List of sections with heading and content"
                },
                "sources": {"type": "array", "items": {"type": "string"}, "description": "List of source URLs or citations"}
            },
            "required": ["filename", "title", "summary", "sections", "sources"]
        }
    }
}

_BASE_TOOL_DEFINITIONS = [
    {
        "name": "web_search",
        "description": "Search the web for up-to-date information. Use when you need facts or data.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search query"},
                "max_results": {"type": "integer", "description": "Number of results (default 5)", "default": 5}
            },
            "required": ["query"]
        }
    },
    {
        "name": "read_file",
        "description": "Read content of a previously saved report.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Path to the file"}
            },
            "required": ["filepath"]
        }
    },
    {
        "name": "merge_results",
        "description": "Merge multiple search result blocks into one formatted block before writing the report.",
        "input_schema": {
            "type": "object",
            "properties": {
                "results": {"type": "array", "items": {"type": "string"}, "description": "List of search strings to merge"}
            },
            "required": ["results"]
        }
    }
]

# Default full definitions (docx)
TOOL_DEFINITIONS = _BASE_TOOL_DEFINITIONS + [_WRITER_TOOL["docx"]]


def get_tool_definitions(output_format: str = "docx") -> list:
    """Return tool definitions with the correct writer for the given format."""
    fmt = output_format if output_format in _WRITER_TOOL else "docx"
    return _BASE_TOOL_DEFINITIONS + [_WRITER_TOOL[fmt]]


TOOL_MAP = {
    "web_search": web_search,
    "docx_writer": docx_writer,
    "txt_writer": txt_writer,
    "read_file": read_file,
    "merge_results": merge_results
}


def execute_tool(tool_name: str, tool_input: dict) -> str:
    if tool_name not in TOOL_MAP:
        return f"Unknown tool: {tool_name}"
    return TOOL_MAP[tool_name](**tool_input)
